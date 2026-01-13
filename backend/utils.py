from __future__ import annotations

import io
import re
import uuid
import time
import os
import tempfile
import json
import base64
import requests
import urllib.parse
from dataclasses import dataclass
from pathlib import Path
from typing import Iterator

import pdfplumber
import fitz
from docx import Document

from .llm import call_vl_llm

# Baidu OCR credentials: prefer environment variables for safety
BAIDU_API_KEY = os.getenv("BAIDU_API_KEY")
BAIDU_SECRET_KEY = os.getenv("BAIDU_SECRET_KEY")

if not BAIDU_API_KEY or not BAIDU_SECRET_KEY:
    # Try one more time in case load_dotenv was delayed
    from dotenv import load_dotenv
    load_dotenv()
    BAIDU_API_KEY = os.getenv("BAIDU_API_KEY")
    BAIDU_SECRET_KEY = os.getenv("BAIDU_SECRET_KEY")

_baidu_token_cache: dict = {"token": None, "expires_at": 0}


def _get_baidu_access_token() -> str | None:
    now = time.time()
    token = _baidu_token_cache.get("token")
    if token and now < _baidu_token_cache.get("expires_at", 0) - 10:
        return token

    if not BAIDU_API_KEY or not BAIDU_SECRET_KEY:
        print("ERROR: BAIDU_API_KEY or BAIDU_SECRET_KEY not set.")
        return None

    url = (
        f"https://aip.baidubce.com/oauth/2.0/token?grant_type=client_credentials&client_id={urllib.parse.quote_plus(BAIDU_API_KEY)}&client_secret={urllib.parse.quote_plus(BAIDU_SECRET_KEY)}"
    )
    try:
        resp = requests.post(url, headers={"Content-Type": "application/json"}, timeout=10)
        data = resp.json()
        if "error" in data:
            print(f"Baidu Token Error: {data.get('error_description')}")
            return None
        token = data.get("access_token")
        expires_in = int(data.get("expires_in", 2592000))
        if token:
            _baidu_token_cache["token"] = token
            _baidu_token_cache["expires_at"] = now + expires_in
            return token
    except Exception as e:
        print(f"Baidu Access Token Exception: {e}")
        return None
    return None


def _call_baidu_seal_api(image_b64: str) -> dict | None:
    token = _get_baidu_access_token()
    if not token:
        return None
    url = f"https://aip.baidubce.com/rest/2.0/ocr/v1/seal?access_token={token}"
    headers = {"Content-Type": "application/x-www-form-urlencoded"}
    data = {"image": image_b64}
    try:
        resp = requests.post(url, data=data, headers=headers, timeout=15)
        return resp.json()
    except Exception:
        return None


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt"}


@dataclass
class PageContent:
    page_number: int
    text: str


@dataclass
class ParsedDocument:
    filename: str
    pages: list[PageContent]


COMPANY_PATTERN = re.compile(r"([\w\u4e00-\u9fff]{2,60}?(?:公司|集团|有限责任公司|股份有限公司))")


SIGNATURE_KEYWORDS = ["签名", "签字", "公司章", "盖章", "印章"]


def is_supported(filename: str) -> bool:
    return Path(filename).suffix.lower() in SUPPORTED_EXTENSIONS

def detect_signature_or_seal(text: str) -> bool:
    return any(keyword in text for keyword in SIGNATURE_KEYWORDS)


def parse_pdf(filename: str, content: bytes) -> ParsedDocument:
    pages: list[PageContent] = []
    with pdfplumber.open(io.BytesIO(content)) as pdf:
        for index, page in enumerate(pdf.pages, start=1):
            text = page.extract_text() or ""
            pages.append(PageContent(page_number=index, text=text))
    return ParsedDocument(filename=filename, pages=pages)


def parse_docx(filename: str, content: bytes) -> ParsedDocument:
    doc = Document(io.BytesIO(content))
    full_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    pages = [PageContent(page_number=1, text=full_text)]
    return ParsedDocument(filename=filename, pages=pages)


def parse_txt(filename: str, content: bytes) -> ParsedDocument:
    text = content.decode("utf-8", errors="ignore")
    return ParsedDocument(filename=filename, pages=[PageContent(page_number=1, text=text)])


def parse_document(filename: str, content: bytes) -> ParsedDocument:
    suffix = Path(filename).suffix.lower()
    if suffix == ".pdf":
        return parse_pdf(filename, content)
    if suffix == ".docx":
        return parse_docx(filename, content)
    if suffix == ".txt":
        return parse_txt(filename, content)
    raise ValueError(f"Unsupported file type: {suffix}")


def summarize_text(text: str, max_len: int = 240) -> str:
    cleaned = re.sub(r"\s+", " ", text).strip()
    if not cleaned:
        return "无可用文本用于总结。"
    if len(cleaned) <= max_len:
        return cleaned
    return f"{cleaned[:max_len]}..."


def iter_pages_for_summary(pages: list[PageContent]) -> Iterator[tuple[int, str]]:
    for page in pages:
        yield page.page_number, summarize_text(page.text)


def render_pdf_page_to_jpg(content: bytes, page_number: int, output_path: Path) -> Path:
    doc = fitz.open(stream=content, filetype="pdf")
    page = doc.load_page(page_number - 1)
    pix = page.get_pixmap(dpi=200)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    pix.save(str(output_path))
    doc.close()
    return output_path


def get_pdf_page_count(content: bytes) -> int:
    """Quickly get the number of pages in a PDF without full parsing."""
    try:
        with fitz.open(stream=content, filetype="pdf") as doc:
            return doc.page_count
    except Exception:
        return 0

def detect_signature_in_image(image_bytes: bytes) -> tuple[bool, str | None]:
    """Use Baidu Seal API to detect signature/seal in an image."""
    try:
        img_b64 = base64.b64encode(image_bytes).decode()
        resp = _call_baidu_seal_api(img_b64)
        if not resp:
            return False, "Failed to call detection API"
        
        if isinstance(resp, dict):
            if resp.get("error_code"):
                return False, f"API Error: {resp.get('error_msg')}"
            
            if resp.get("result_num", 0) > 0:
                return True, "检测到印章/签名。"
            
            # Fallback to OCR text
            words_result = resp.get("words_result")
            if words_result:
                for item in words_result:
                    if detect_signature_or_seal(item.get("words", "")):
                        return True, "文字中包含签名/盖章关键词。"
    except Exception as e:
        return False, f"Internal Error: {str(e)}"
    
    return False, "未检测到印章或签名。"



def detect_signature_in_pdf(pdf_bytes: bytes, num_pages: int) -> tuple[bool, str | None]:
    """Use VLM to inspect PDF pages for signatures/seals.

    Returns (found: bool, note: Optional[str]). 
    Optimized: Checks first 3 and last 5 pages first, then the rest.
    """
    if not pdf_bytes or num_pages is None or num_pages == 0:
        return False, None

    # Prioritize finding seals near the front (cover/intro) or back (declaration/assurance)
    # Most ESG reports have seals on the last few pages.
    pages_to_check = []
    if num_pages <= 10:
        pages_to_check = list(range(1, num_pages + 1))
    else:
        # First 3
        pages_to_check.extend([1, 2, 3])
        # Last 5
        last_pages = [num_pages - i for i in range(5)]
        for lp in last_pages:
            if lp > 3:
                pages_to_check.append(lp)
        # Remaining pages
        checked = set(pages_to_check)
        for p in range(1, num_pages + 1):
            if p not in checked:
                pages_to_check.append(p)

    for p in pages_to_check:
        tmp_path = None
        try:
            # create a secure temporary file for the rendered page
            tf = tempfile.NamedTemporaryFile(suffix=".jpg", delete=False)
            tmp_path = Path(tf.name)
            tf.close()
            
            image_path = render_pdf_page_to_jpg(pdf_bytes, p, tmp_path)
            img_bytes = Path(image_path).read_bytes()
            img_b64 = base64.b64encode(img_bytes).decode()
            resp = _call_baidu_seal_api(img_b64)
            
            if not resp:
                continue
            
            if isinstance(resp, dict):
                error_code = resp.get("error_code")
                if error_code:
                    continue
                
                # Check for successful detection
                # result = resp.get("result")
                words_result = resp.get("words_result") # fallback for some OCR types
                
                if resp.get("result_num", 0) > 0:
                    return True, f"在第 {p} 页检测到印章/签名。"
                elif words_result and len(words_result) > 0:
                    # Fallback: check if any recognized text contains signature keywords
                    for item in words_result:
                        words = item.get("words", "")
                        if detect_signature_or_seal(words):
                            return True, f"在第 {p} 页检测到印章/签名。"
        except Exception:
            continue
        finally:
            if tmp_path and tmp_path.exists():
                try:
                    tmp_path.unlink()
                except:
                    pass

    return False, None


def detect_signature_with_vl_model(pdf_bytes: bytes, num_pages: int) -> tuple[bool, str | None]:
    """Use VL model to inspect PDF pages for handwritten signatures.

    Returns (found: bool, note: Optional[str]). 
    Optimized: Checks first 3 and last 5 pages first, then the rest.
    """
    if not pdf_bytes or num_pages is None or num_pages == 0:
        return False, None

    # Prioritize finding signatures near the front (cover/intro) or back (declaration/assurance)
    pages_to_check = []
    if num_pages <= 10:
        pages_to_check = list(range(1, num_pages + 1))
    else:
        # First 3
        pages_to_check.extend([1, 2, 3])
        # Last 5
        last_pages = [num_pages - i for i in range(5)]
        for lp in last_pages:
            if lp > 3:
                pages_to_check.append(lp)
        # Remaining pages
        checked = set(pages_to_check)
        for p in range(1, num_pages + 1):
            if p not in checked:
                pages_to_check.append(p)

    prompt = "这张图片中是否有手写签名？请回答yes或no，不要简要说明。"

    for p in pages_to_check:
        tmp_path = None
        try:
            # create a secure temporary file for the rendered page
            tf = tempfile.NamedTemporaryFile(suffix=".jpg", delete=False)
            tmp_path = Path(tf.name)
            tf.close()
            
            image_path = render_pdf_page_to_jpg(pdf_bytes, p, tmp_path)
            img_bytes = Path(image_path).read_bytes()
            img_b64 = base64.b64encode(img_bytes).decode()
            
            response = call_vl_llm(img_b64, prompt)
            
            # Parse response: check if contains "yes" (assuming English response)
            if "yes" in response.lower():
                return True, f"在第 {p} 页检测到手写签名。"
        except Exception:
            continue
        finally:
            if tmp_path and tmp_path.exists():
                try:
                    tmp_path.unlink()
                except:
                    pass

    return False, None
